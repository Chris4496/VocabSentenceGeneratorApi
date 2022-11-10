from docx import Document
from docx.shared import Pt
from fpdf import FPDF, XPos, YPos
import os
import tempfile


def generateTempTextFile(examples, answers):
    # add numbers to answers
    answers = [f"{i+1}. {answers[i]}" for i in range(len(answers))]

    fd, path = tempfile.mkstemp(suffix=".txt")
    with os.fdopen(fd, "w") as f:
        for i in range(len(examples)):
            f.write(f"{i+1}. {examples[i]}\n")

        f.write("\n")
        f.write(f'{"-" * 20}\n')
        f.write(f"Answers:\n")
        f.write(" | ".join(answers))

    return path


def generateTempDocxFile(examples, answers):
    # add numbers to answers
    answers = [f"{i+1}. {answers[i]}" for i in range(len(answers))]

    fd, path = tempfile.mkstemp(suffix=".docx")
    with os.fdopen(fd, "w") as f:
        document = Document()

        # set font to Calibri
        font = document.styles["Normal"].font
        font.name = "Calibri"
        font.size = Pt(13)

        for i in range(len(examples)):
            document.add_paragraph(examples[i], style="List Number")

        document.add_paragraph("-" * 20)
        document.add_paragraph("Answers:")
        document.add_paragraph(" | ".join(answers))

        document.save(path)

    return path


def generateTempPDFFile(examples, answers):
    # add numbers to answers
    answers = [f"{i+1}. {answers[i]}" for i in range(len(answers))]

    fd, path = tempfile.mkstemp(suffix=".pdf")
    with os.fdopen(fd, "w") as f:
        pdf = FPDF()
        pdf.add_page()
        pdf.add_font("Calibri", "", "fonts/calibri.ttf")
        pdf.set_font("Calibri", size=13)

        for i in range(len(examples)):
            pdf.cell(
                200, 10, txt=f"{i+1}. {examples[i]}", new_x=XPos.LEFT, new_y=YPos.NEXT, align="L")

        pdf.cell(200, 10, txt="-" * 20, new_x=XPos.LEFT,
                 new_y=YPos.NEXT, align="L")
        pdf.cell(200, 10, txt="Answers:", new_x=XPos.LEFT,
                 new_y=YPos.NEXT, align="L")

        for answer in answers:
            pdf.cell(200, 10, txt=answer, new_x=XPos.LEFT,
                     new_y=YPos.NEXT, align="L")

        pdf.output(path)

    return path


if __name__ == "__main__":
    examples = ["This is an example", "This is another example", "This is a third example", "This is a fourth example", "This is a fifth example", "This is a sixth example", "This is a seventh example", "This is an eighth example", "This is a ninth example", "This is a tenth example", "This is an eleventh example", "This is a twelfth example", "This is a thirteenth example", "This is a fourteenth example", "This is a fifteenth example", "This is a sixteenth example", "This is a seventeenth example", "This is an eighteenth example", "This is a nineteenth example", "This is a twentieth example", "This is a twenty-first example", "This is a twenty-second example", "This is a twenty-third example", "This is a twenty-fourth example", "This is a twenty-fifth example", "This is a twenty-sixth example", "This is a twenty-seventh example", "This is a twenty-eighth example", "This is a twenty-ninth example", "This is a thirtieth example", "This is a thirty-first example", "This is a thirty-second example", "This is a thirty-third example", "This is a thirty-fourth example", "This is a thirty-fifth example",
                "This is a thirty-sixth example", "This is a thirty-seventh example", "This is a thirty-eighth example", "This is a thirty-ninth example", "This is a fortieth example", "This is a forty-first example", "This is a forty-second example", "This is a forty-third example", "This is a forty-fourth example", "This is a forty-fifth example", "This is a forty-sixth example", "This is a forty-seventh example", "This is a forty-eighth example", "This is a forty-ninth example", "This is a fiftieth example", "This is a fifty-first example", "This is a fifty-second example", "This is a fifty-third example", "This is a fifty-fourth example", "This is a fifty-fifth example", "This is a fifty-sixth example", "This is a fifty-seventh example", "This is a fifty-eighth example", "This is a fifty-ninth example", "This is a sixtieth example", "This is a sixty-first example", "This is a sixty-second example", "This is a sixty-third example", "This is a sixty-fourth example", "This is a sixty-fifth example", "This is a sixty-sixth example", "This is a sixty-seventh example", "This is a sixty-eighth example"]
    answers = ["This is an answer", "This is another answer", "This is a third answer", "This is a fourth answer", "This is a fifth answer", "This is a sixth answer", "This is a seventh answer", "This is an eighth answer", "This is a ninth answer", "This is a tenth answer", "This is an eleventh answer", "This is a twelfth answer", "This is a thirteenth answer", "This is a fourteenth answer", "This is a fifteenth answer", "This is a sixteenth answer", "This is a seventeenth answer", "This is an eighteenth answer", "This is a nineteenth answer", "This is a twentieth answer", "This is a twenty-first answer", "This is a twenty-second answer", "This is a twenty-third answer", "This is a twenty-fourth answer", "This is a twenty-fifth answer", "This is a twenty-sixth answer", "This is a twenty-seventh answer", "This is a twenty-eighth answer", "This is a twenty-ninth answer", "This is a thirtieth answer", "This is a thirty-first answer", "This is a thirty-second answer", "This is a thirty-third answer", "This is a thirty-fourth answer", "This is a thirty-fifth answer",
               "This is a thirty-sixth answer", "This is a thirty-seventh answer", "This is a thirty-eighth answer", "This is a thirty-ninth answer", "This is a fortieth answer", "This is a forty-first answer", "This is a forty-second answer", "This is a forty-third answer", "This is a forty-fourth answer", "This is a forty-fifth answer", "This is a forty-sixth answer", "This is a forty-seventh answer", "This is a forty-eighth answer", "This is a forty-ninth answer", "This is a fiftieth answer", "This is a fifty-first answer", "This is a fifty-second answer", "This is a fifty-third answer", "This is a fifty-fourth answer", "This is a fifty-fifth answer", "This is a fifty-sixth answer", "This is a fifty-seventh answer", "This is a fifty-eighth answer", "This is a fifty-ninth answer", "This is a sixtieth answer", "This is a sixty-first answer", "This is a sixty-second answer", "This is a sixty-third answer", "This is a sixty-fourth answer", "This is a sixty-fifth answer", "This is a sixty-sixth answer", "This is a sixty-seventh answer", "This is a sixty-eighth answer"]
    print(generateTempTextFile(examples, answers))
